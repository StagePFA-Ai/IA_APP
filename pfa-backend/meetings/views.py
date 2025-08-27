# views.py
# =============================================================================
# MeetingAI — Vues Django (corrigées & commentées)
# Politique d'accès :
#   - Utilisateur : accès à ses données (créateur/participant).
#   - Admin (staff/superuser) : statistiques globales uniquement (pas de contenu).
# =============================================================================

from __future__ import annotations
import os
from tempfile import NamedTemporaryFile
from datetime import date as ddate, datetime, time as dtime

from django.conf import settings
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout, get_user_model
from django.contrib.auth.decorators import login_required
from django.db import IntegrityError, transaction
from django.db.models import Q, Sum, Count
from django.db.models.functions import TruncMonth
from django.http import (
    FileResponse,
    HttpResponse,
    HttpResponseBadRequest,
    HttpResponseForbidden,
    JsonResponse,
)
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import NoReverseMatch, reverse
from django.utils import timezone
from django.utils.dateparse import parse_date
from django.core.paginator import Paginator

from docx import Document
from docx.shared import Inches

from .models import Reunion, Audio, Transcription, Resume, Rapport

User = get_user_model()

# =============================================================================
# Constantes statut — harmonisées (adaptez si vos choices diffèrent)
# =============================================================================
STATUS_LABELS = {
    "planifier": "Planifié",
    "en_cours": "En cours",
    "terminer": "Terminé",
    "annuler": "Annulé",
    "reporter": "Reporté",
}

STATUS_COLORS = {
    "planifier": "#2563eb",  # bleu
    "en_cours": "#10b981",   # vert
    "terminer": "#6b7280",   # gris
    "annuler": "#ef4444",    # rouge
    "reporter": "#f59e0b",   # orange
}

# =============================================================================
# Helpers rôles & permissions
# =============================================================================
def is_admin(user) -> bool:
    """Admin = staff OU superuser."""
    return bool(user and (user.is_staff or user.is_superuser))


def forbid_admin_on_meetings(request) -> bool:
    """
    Bloque l'accès 'réunions/calendrier' aux administrateurs :
    l'admin n'a qu'une vision statistique, pas de contenu.
    """
    if is_admin(request.user):
        messages.info(
            request,
            "Accès réservé aux utilisateurs. En tant qu'administrateur, vous avez une vue statistiques uniquement.",
        )
        return True
    return False


def _today_local():
    return timezone.localdate()


def _parse_hhmm(s: str | None):
    """'HH:MM' -> datetime.time | None"""
    try:
        h, m = (s or "").split(":")
        return dtime(int(h), int(m))
    except Exception:
        return None


def _scope_user(user: User) -> Q: # type: ignore
    """Réunions visibles par l'utilisateur : créateur OU participant."""
    return Q(utilisateur=user) | Q(participants=user)


def _can_view_meeting(reunion: Reunion, user: User) -> bool: # type: ignore
    """
    Droits d'accès au contenu d'une réunion :
    - créateur
    - OU participant

    ⚠️ Pas de passe-droit admin : l'admin ne lit pas le contenu des réunions d'autrui.
    """
    return reunion.utilisateur_id == user.id or reunion.participants.filter(id=user.id).exists()


def _can_edit(reunion: Reunion, user: User) -> bool: # type: ignore
    """Seul le créateur peut modifier :
       - si réunion aujourd'hui ou future
       - si 'reporter' (pour replanifier)
    """
    if reunion.utilisateur_id != user.id:
        return False
    if reunion.status == "reporter":
        return True
    return reunion.date_r >= _today_local()


def _can_start(reunion: Reunion) -> bool:
    """Démarrage autorisé uniquement le jour J si statut 'planifier'."""
    return reunion.status == "planifier" and reunion.date_r == _today_local()


def _can_resume(reunion: Reunion) -> bool:
    """Rejoindre la transcription si réunion en cours le jour J."""
    return reunion.status == "en_cours" and reunion.date_r == _today_local()


def _autopostpone_overdue(user: User): # type: ignore
    """Marque 'reporter' les réunions du user dépassées restées 'planifier'."""
    today = _today_local()
    Reunion.objects.filter(utilisateur=user, status="planifier", date_r__lt=today).update(status="reporter")


# =============================================================================
# Accueil / Auth
# =============================================================================
#la 1er page de l'application 
def home(request):
    return render(request, "HTML/home.html")

#page login 
def login_view(request):
    """Login simple avec message d'erreur et redirection vers dashboard."""
    if request.method == "POST":
        username = (request.POST.get("username") or "").strip()
        password = request.POST.get("password") or ""
        user = authenticate(request, username=username, password=password)
        if user and user.is_active:
            login(request, user)
            return redirect(request.GET.get("next") or "dashboard")
        messages.error(request, "Identifiants incorrects")
    return render(request, "HTML/login.html")
#page logout
def logout_view(request):
    """Logout + retour page de connexion."""
    logout(request)
    return redirect("login")


# =============================================================================
# Dashboard : stats globales (admin) vs stats personnelles (utilisateur)
# =============================================================================
#view dashboard
@login_required
def dashboard(request):
    user = request.user
    today = _today_local()

    # --- ADMIN : statistiques globales uniquement (aucune liste détaillée) ---
    if is_admin(user):
        all_meetings = Reunion.objects.all()
        total_meetings = all_meetings.count()
        total_seconds = Audio.objects.aggregate(s=Sum("duree")).get("s") or 0
        heures = round(total_seconds / 3600, 1)
        resumes_count = Resume.objects.count()

        # Histogramme global par mois
        meetings_by_month = (
            Reunion.objects
            .annotate(month=TruncMonth('date_r'))
            .values('month')
            .annotate(count=Count('id'))
            .order_by('month')
        )
        chart1_labels = [m['month'].strftime("%b %Y") for m in meetings_by_month]
        chart1_data = [m['count'] for m in meetings_by_month]

        ctx = {
            "kpis": {
                "total_meetings": total_meetings,
                "heures": heures,
                "Résumés": resumes_count,
                "actions": 0,  # pas d’actions ciblées en mode admin
            },
            "chart1_labels": chart1_labels,
            "chart1_data": chart1_data,
            "recent_meetings": [],  # on n’affiche pas de contenu détaillé à l’admin
            "admin_mode": True,
        }
        return render(request, "HTML/dashboard.html", ctx)

    # --- UTILISATEUR : stats personnelles filtrées à son scope ---
    meetings_qs = Reunion.objects.filter(_scope_user(user)).distinct()

    total_meetings = meetings_qs.count()
    total_seconds = Audio.objects.filter(reunion__in=meetings_qs).aggregate(s=Sum("duree")).get("s") or 0
    heures = round(total_seconds / 3600, 1)
    resumes_count = Resume.objects.filter(transcription__reunion__in=meetings_qs).count()

    actions = meetings_qs.filter(status__in=["planifier", "en_cours"], date_r__gte=today).count()

    # Histogramme personnel par mois
    meetings_by_month = (
        meetings_qs.filter(_scope_user(user))
        .annotate(month=TruncMonth('date_r'))
        .values('month')
        .annotate(count=Count('id'))
        .order_by('month')
    )
    chart1_labels = [m['month'].strftime("%b %Y") for m in meetings_by_month]
    chart1_data = [m['count'] for m in meetings_by_month]
    recent_meetings = list(meetings_qs.order_by("-date_r", "-heure_r")[:8])

    ctx = {
        "kpis": {
            "total_meetings": total_meetings,
            "heures": heures,
            "Résumés": resumes_count,
            "actions": actions,
        },
        "chart1_labels": chart1_labels,
        "chart1_data": chart1_data,
        "recent_meetings": recent_meetings,
        "admin_mode": False,
    }
    return render(request, "HTML/dashboard.html", ctx)
# =============================================================================
# Calendrier — interdit aux admins (stats only)
# =============================================================================
@login_required
def calendar(request):
    """Page calendrier (colonne droite préremplie avec aujourd'hui)."""
    if forbid_admin_on_meetings(request):
        return redirect("dashboard")
    selected_date = parse_date(request.GET.get("date") or "") or _today_local()
    meetings_on_date = (
        Reunion.objects
        .filter(_scope_user(request.user), date_r=selected_date)
        .select_related("utilisateur")
        .prefetch_related("participants")
        .order_by("heure_r", "titre")
        .distinct()
    )

    _autopostpone_overdue(request.user)

    return render(request, "HTML/calendar.html", {
        "selected_date": selected_date,
        "meetings_on_date": meetings_on_date,
    })


@login_required
def calendar_events(request):
    """Événements pour FullCalendar (toutes réunions visibles de l'utilisateur)."""
    if forbid_admin_on_meetings(request):
        return JsonResponse({"detail": "Admin: stats only."}, status=403)

    qs = (
        Reunion.objects
        .filter(_scope_user(request.user))
        .select_related("utilisateur")
        .distinct()
    )
    events = [{
        "id": r.id,
        "title": r.titre,
        "start": f"{r.date_r.isoformat()}T{r.heure_r.strftime('%H:%M')}",
        "url": reverse("view_meeting", args=[r.id]),
        "color": STATUS_COLORS.get(r.status, "#2563eb"),
    } for r in qs]
    return JsonResponse(events, safe=False)


@login_required
def calendar_day(request):
    """Réunions d’un jour (pour remplir la colonne de droite)."""
    if forbid_admin_on_meetings(request):
        return JsonResponse({"detail": "Admin: stats only."}, status=403)

    try:
        day_str = request.GET.get("date", "")
        d = parse_date(day_str) or _today_local()
    except Exception:
        return HttpResponseBadRequest("date invalide")

    _autopostpone_overdue(request.user)

    qs = (
        Reunion.objects
        .filter(_scope_user(request.user), date_r=d)
        .select_related("utilisateur")
        .prefetch_related("participants")
        .order_by("heure_r", "titre")
        .distinct()
    )

    payload = {"date": d.isoformat(), "meetings": []}
    for m in qs:
        can_start = _can_start(m)
        can_resume = _can_resume(m)

        if can_resume:
            hint = "Réunion en cours — vous pouvez la rejoindre."
        elif m.status == "reporter":
            hint = "Réunion dépassée non démarrée : replanifie-la."
        elif not can_start and m.date_r == _today_local():
            hint = "Le démarrage n’est possible que si le statut est ‘Planifié’."
        elif not can_start and m.date_r != _today_local():
            hint = "Le démarrage est autorisé uniquement le jour J."
        else:
            hint = ""

        payload["meetings"].append({
            "id": m.id,
            "titre": m.titre,
            "date": m.date_r.isoformat(),
            "heure": m.heure_r.strftime("%H:%M"),
            "status": m.status,
            "status_label": STATUS_LABELS.get(m.status, m.status),
            "can_start": can_start,
            "can_resume": can_resume,
            "resume_url": reverse("transcription", args=[m.id]) if can_resume else "",
            "hint": hint,
        })
    return JsonResponse(payload)


@login_required
def calendar_create(request):
    """Créer une réunion (modale ‘Nouvelle réunion’)."""
    if forbid_admin_on_meetings(request):
        return JsonResponse({"ok": False, "error": "Admin: stats only."}, status=403)

    if request.method != "POST":
        return HttpResponseBadRequest("Méthode non supportée")

    titre = (request.POST.get("titre") or "").strip()
    d = parse_date(request.POST.get("date") or "")
    t = _parse_hhmm(request.POST.get("heure") or "")

    if not titre or not d or not t:
        return JsonResponse({"ok": False, "error": "Champs manquants"}, status=400)

    r = Reunion.objects.create(
        titre=titre,
        date_r=d,
        heure_r=t,
        status="planifier",
        utilisateur=request.user,
    )
    return JsonResponse({"ok": True, "id": r.id})


@login_required
def calendar_meeting_info(request, pk: int):
    """Infos pour la modale 'Gérer' (droits + démarrer/rejoindre)."""
    if forbid_admin_on_meetings(request):
        return JsonResponse({"detail": "Admin: stats only."}, status=403)

    try:
        m = (
            Reunion.objects
            .select_related("utilisateur")
            .prefetch_related("participants")
            .get(Q(pk=pk) & _scope_user(request.user))
        )
    except Reunion.DoesNotExist:
        return HttpResponseBadRequest("Réunion introuvable ou non autorisée")

    # Auto-report si planifiée & dépassée
    if m.status == "planifier" and m.date_r < _today_local():
        m.status = "reporter"
        m.save(update_fields=["status"])

    can_start = _can_start(m)
    can_resume = _can_resume(m)
    can_edit = _can_edit(m, request.user)

    # Participants (tous les users affichés — adaptez au besoin)
    selected_ids = set(m.participants.values_list("id", flat=True))
    participants_all = []
    for u in User.objects.order_by("username").values("id", "username", "first_name", "last_name"):
        label = (f"{u['first_name']} {u['last_name']}".strip()) or u["username"]
        participants_all.append({"id": u["id"], "label": label, "selected": u["id"] in selected_ids})

    if can_resume:
        hint = "Réunion en cours — vous pouvez la rejoindre."
    elif m.status == "reporter":
        hint = "Cette réunion a été reportée car elle a dépassé sa date sans démarrage. Replanifie-la."
    elif m.date_r == _today_local() and not can_start:
        hint = "Le démarrage n’est possible que si le statut est ‘Planifié’."
    elif m.date_r != _today_local() and not can_start:
        hint = "Le démarrage est autorisé uniquement le jour J."
    else:
        hint = ""

    data = {
        "id": m.id,
        "titre": m.titre,
        "date": m.date_r.isoformat(),
        "heure": m.heure_r.strftime("%H:%M"),
        "status": m.status,
        "status_label": STATUS_LABELS.get(m.status, m.status),
        "can_start": can_start,
        "can_edit": can_edit,
        "can_resume": can_resume,
        "resume_url": reverse("transcription", args=[m.id]) if can_resume else "",
        "participants_all": participants_all,
        "hint": hint,
    }
    return JsonResponse(data)


@login_required
def calendar_meeting_update(request):
    """Met à jour date/heure/participants (créateur uniquement)."""
    if forbid_admin_on_meetings(request):
        return JsonResponse({"ok": False, "error": "Admin: stats only."}, status=403)

    if request.method != "POST":
        return HttpResponseBadRequest("Méthode invalide")

    pk = request.POST.get("id")
    try:
        m = Reunion.objects.select_related("utilisateur").get(pk=pk)
    except Reunion.DoesNotExist:
        return JsonResponse({"ok": False, "error": "Réunion introuvable"}, status=400)

    if not _can_edit(m, request.user):
        return JsonResponse({"ok": False, "error": "Modification non autorisée"}, status=403)

    # maj date/heure
    d = parse_date(request.POST.get("date") or "")
    t = _parse_hhmm(request.POST.get("heure") or "")
    if not d or not t:
        return JsonResponse({"ok": False, "error": "Date/heure invalides"}, status=400)
    m.date_r, m.heure_r = d, t

    # participants
    ids = request.POST.getlist("participants")
    if ids:
        m.participants.set(User.objects.filter(id__in=ids))
    else:
        m.participants.clear()

    # si 'reporter' et replanifiée aujourd’hui/futur → redevient planifiée
    if m.status == "reporter" and m.date_r >= _today_local():
        m.status = "planifier"

    m.save()
    return JsonResponse({"ok": True})


@login_required
def calendar_start(request):
    """Démarrage d’une réunion (jour J & statut planifier)."""
    if forbid_admin_on_meetings(request):
        return JsonResponse({"ok": False, "error": "Admin: stats only."}, status=403)

    if request.method != "POST":
        return HttpResponseBadRequest("Méthode invalide")

    pk = request.POST.get("id")
    try:
        m = Reunion.objects.select_related("utilisateur").get(pk=pk)
    except Reunion.DoesNotExist:
        return JsonResponse({"ok": False, "error": "Réunion introuvable"}, status=400)

    if not _can_start(m):
        return JsonResponse(
            {"ok": False, "error": "Le démarrage est autorisé uniquement le jour J pour une réunion planifiée."},
            status=403,
        )

    if m.status != "en_cours":
        m.status = "en_cours"
        m.save(update_fields=["status"])

    return JsonResponse({"ok": True, "redirect": reverse("transcription", args=[m.id])})


# =============================================================================
# Détail / Historique (interdit aux admins)
# =============================================================================
@login_required
def view_meeting(request, meeting_id: int):
    """Fiche détaillée d’une réunion (créateur/participant uniquement)."""
    if forbid_admin_on_meetings(request):
        return HttpResponseForbidden("Accès interdit.")

    m = get_object_or_404(Reunion.objects.select_related("utilisateur"), id=meeting_id)
    if not _can_view_meeting(m, request.user):
        return HttpResponseForbidden("Accès interdit.")

    audios = m.audios.all().order_by("-id")
    total_seconds = audios.aggregate(s=Sum("duree")).get("s") or 0
    total_hours = round(total_seconds / 3600, 2)

    transcription = getattr(m, "transcription", None)
    resume = getattr(transcription, "resume", None) if transcription else None
    rapport = getattr(resume, "rapport", None) if resume else None

    return render(request, "HTML/meeting_detail.html", {
        "meeting": m,
        "is_owner": (m.utilisateur_id == request.user.id),
        "participants": m.participants.all(),
        "audios": audios,
        "total_audio_seconds": total_seconds,
        "total_audio_hours": total_hours,
        "transcription": transcription,
        "resume": resume,
        "rapport": rapport,
    })


@login_required
def meetings_page(request):
    """
    Liste des réunions de l'utilisateur (créées par lui ou où il est participant),
    avec recherche plein texte (titre, transcription, résumé), filtres statut / dates,
    et pagination.
    """
    if forbid_admin_on_meetings(request):
        return redirect("dashboard")

    user = request.user

    qs = (
        Reunion.objects
        .filter(_scope_user(user))
        .distinct()
    )

    # --- Filtres GET ---
    q = (request.GET.get("q") or "").strip()
    status = (request.GET.get("status") or "").strip()  # planifier, en_cours, terminer, annuler, reporter
    dfrom = parse_date(request.GET.get("date_from") or "")
    dto = parse_date(request.GET.get("date_to") or "")

    if q:
        qs = qs.filter(
            Q(titre__icontains=q) |
            Q(transcription__text_transcrit__icontains=q) |
            Q(transcription__resume__text_resume__icontains=q)
        )

    if status in {"planifier", "en_cours", "terminer", "annuler", "reporter"}:
        qs = qs.filter(status=status)

    if dfrom:
        qs = qs.filter(date_r__gte=dfrom)
    if dto:
        qs = qs.filter(date_r__lte=dto)

    qs = (
        qs.select_related("utilisateur", "transcription", "transcription__resume")
          .prefetch_related("participants", "audios")
          .order_by("-date_r", "-heure_r", "-id")
    )

    paginator = Paginator(qs, 8)  # 8 cartes par page
    page_obj = paginator.get_page(request.GET.get("page"))

    ctx = {
        "meetings": page_obj.object_list,
        "page_obj": page_obj,
        "q": q,
        "status": status,
        "date_from": dfrom.isoformat() if dfrom else "",
        "date_to": dto.isoformat() if dto else "",
    }
    return render(request, "HTML/meetings.html", ctx)


# =============================================================================
# Paramètres (session) — accessible à tous (admin inclus si nécessaire)
# =============================================================================
@login_required
def settings_page(request):
    """
    Paramètres en SESSION (aucun changement de models).
    Clés session:
      user_settings = {
        auto_record: bool,
        recording_consent: bool,
        doc_notes: str,
      }
    """
    sess = request.session.setdefault("user_settings", {})

    if request.method == "POST":
        def asbool(v): return str(v).lower() in {"1", "true", "on", "yes", "oui"}
        sess["auto_record"] = asbool(request.POST.get("auto_record"))
        sess["recording_consent"] = asbool(request.POST.get("recording_consent"))

        request.session["user_settings"] = sess
        request.session.modified = True
        messages.success(request, "Paramètres enregistrés.")
        return redirect("settings")

    ctx = {
        "settings": {
            "auto_record": bool(sess.get("auto_record", False)),
            "recording_consent": bool(sess.get("recording_consent", True)),
        },
        "user_info": request.user,
        # Texte affiché
        "app_description": (
            "MeetingAI est une application locale (on-premise) qui transcrit les réunions "
            "avec Whisper / faster-whisper et génère des résumés multilingues (FR/EN/AR). "
            "Toutes les données restent dans l’infrastructure de l’entreprise."
        ),
        "arch_points": [
            "Frontend Django + Bootstrap (Dashboard, Réunions, Calendrier, Paramètres).",
            "Capture micro navigateur (WebRTC) → backend ASGI.",
            "Transcodage en PCM 16 kHz mono → Transcription (Whisper/faster-whisper).",
            "Résumé multilingue (mT5/LLM) côté backend.",
            "Persistance : Audio, Transcription, Résumé, Rapport (suivant permissions).",
        ],
        # Chemins purement informatifs (lecture seule)
        "media_root": settings.MEDIA_ROOT,
        "audio_upload_path": "uploads/audio/",
        "reports_upload_path": "uploads/rapports/",
        "files_upload_path": "uploads/",
    }
    return render(request, "HTML/settings.html", ctx)


# =============================================================================
# Création / modification réunions (utilisateur uniquement)
# =============================================================================
@login_required
def reunion_form(request, reunion_id=None):
    """Affiche le formulaire : création si None, sinon modification (créateur only)."""
    if forbid_admin_on_meetings(request):
        return redirect("dashboard")

    reunion = None
    if reunion_id:
        reunion = get_object_or_404(Reunion, pk=reunion_id)
        if reunion.utilisateur != request.user:
            messages.error(request, "Vous n'êtes pas autorisé à modifier cette réunion.")
            return redirect("meetings")

    utilisateurs = User.objects.all().order_by("username")  # pour <select>
    return render(request, "HTML/nouvelle_reunion.html", {"reunion": reunion, "utilisateurs": utilisateurs})


@login_required
def creer_reunion(request):
    """Créer une réunion (bouton 'enregistrer' ou 'demarrer')."""
    if forbid_admin_on_meetings(request):
        return redirect("dashboard")

    if request.method == "POST":
        titre = (request.POST.get("titre") or "").strip()
        date_str = request.POST.get("date_r")
        heure_str = request.POST.get("heure_r")
        action = request.POST.get("action")  # "enregistrer" ou "demarrer"

        if not titre or not date_str or not heure_str:
            messages.error(request, "Veuillez renseigner le titre, la date et l’heure.")
            return redirect("creer_reunion")

        try:
            date_r = datetime.strptime(date_str, "%Y-%m-%d").date()
            heure_r = datetime.strptime(heure_str, "%H:%M").time()
        except ValueError:
            messages.error(request, "Format de date/heure invalide.")
            return redirect("creer_reunion")

        # Unicité basique (évite doublon immédiat)
        conflict = Reunion.objects.filter(
            utilisateur=request.user,
            date_r=date_r,
            heure_r=heure_r,
            titre=titre,
        ).exists()
        if conflict:
            messages.error(request, "Une réunion identique existe déjà pour votre compte.")
            return redirect("creer_reunion")

        try:
            with transaction.atomic():
                reunion = Reunion.objects.create(
                    titre=titre,
                    date_r=date_r,
                    heure_r=heure_r,
                    utilisateur=request.user,
                    status="planifier",
                )

                p_ids = request.POST.getlist("participants")
                if p_ids:
                    reunion.participants.set(User.objects.filter(id__in=p_ids))

                if action == "demarrer":
                    reunion.status = "en_cours"
                    reunion.save(update_fields=["status"])
                    messages.success(request, "Réunion démarrée.")
                    return redirect("transcription", reunion_id=reunion.id)

            messages.success(request, "Réunion enregistrée.")
            return redirect("meetings")

        except IntegrityError:
            messages.error(
                request,
                "Conflit détecté : une réunion identique vient d’être créée. Réessayez avec un autre horaire ou un autre titre.",
            )
            return redirect("creer_reunion")

    # GET : afficher le formulaire de création
    utilisateurs = User.objects.exclude(id=request.user.id).order_by("username")
    return render(request, "HTML/nouvelle_reunion.html", {"utilisateurs": utilisateurs})


@login_required
def modifier_reunion(request, reunion_id):
    """POST de modification. 'demarrer' passe la réunion en 'en_cours'."""
    if forbid_admin_on_meetings(request):
        return redirect("dashboard")

    reunion = get_object_or_404(Reunion, pk=reunion_id)
    if reunion.utilisateur != request.user:
        messages.error(request, "Vous n'êtes pas autorisé à modifier cette réunion.")
        return redirect("meetings")

    if request.method != "POST":
        return redirect("reunion_modifier", reunion_id=reunion.id)

    titre = (request.POST.get("titre") or "").strip()
    date_r = request.POST.get("date_r")
    heure_r = request.POST.get("heure_r")
    action = request.POST.get("action")  # 'modifier' ou 'demarrer'

    if not titre or not date_r or not heure_r:
        messages.error(request, "Veuillez remplir tous les champs obligatoires.")
        return redirect("reunion_modifier", reunion_id=reunion.id)

    reunion.titre = titre
    reunion.date_r = date_r
    reunion.heure_r = heure_r

    participants_ids = request.POST.getlist("participants")
    reunion.participants.set(User.objects.filter(id__in=participants_ids))

    if action == "demarrer":
        reunion.status = "en_cours"

    reunion.save()

    if action == "demarrer":
        messages.success(request, "Réunion démarrée.")
        try:
            return redirect("transcription", reunion_id=reunion.id)
        except Exception:
            return redirect(f"/transcription/?reunion_id={reunion.id}")

    messages.success(request, "Réunion mise à jour.")
    return redirect("meetings")


def reunion_nouvelle(request):
    """Alias création."""
    return reunion_form(request, reunion_id=None)


# =============================================================================
# Transcription / Résumé (utilisateur uniquement)
# =============================================================================
@login_required
def transcription_page(request, reunion_id: int):
    if forbid_admin_on_meetings(request):
        return HttpResponseForbidden("Accès interdit.")

    reunion = get_object_or_404(Reunion, pk=reunion_id)
    if not _can_view_meeting(reunion, request.user):
        return HttpResponseForbidden("Accès interdit.")

    transcription = getattr(reunion, "transcription", None)
    resume = getattr(transcription, "resume", None) if transcription else None

    ctx = {
        "reunion": reunion,
        "transcription": transcription,
        "resume": resume,
        "webrtc_ws_url": "/ws/transcription/",
    }
    return render(request, "HTML/transcription.html", ctx)


@login_required
def save_transcription(request, reunion_id: int):
    if forbid_admin_on_meetings(request):
        return JsonResponse({"ok": False, "error": "Admin: stats only."}, status=403)

    if request.method != "POST":
        return HttpResponseForbidden("Méthode non autorisée")

    reunion = get_object_or_404(Reunion, pk=reunion_id)
    if not _can_view_meeting(reunion, request.user):
        return HttpResponseForbidden("Accès interdit.")

    text = (request.POST.get("text") or "").strip()
    summary = (request.POST.get("summary") or "").strip()
    lang = (request.POST.get("lang") or "fr").strip()

    if not text:
        return JsonResponse({"ok": False, "error": "Texte vide"}, status=400)

    # upsert Transcription
    tr, created = Transcription.objects.get_or_create(
        reunion=reunion,
        defaults={"text_transcrit": text, "langue": lang, "heure_date": timezone.now()},
    )
    if not created:
        tr.text_transcrit = text
        tr.langue = lang
        tr.heure_date = timezone.now()
        tr.save()

    # upsert Resume (si summary fourni)
    if summary:
        rs, c2 = Resume.objects.get_or_create(
            transcription=tr, defaults={"text_resume": summary}
        )
        if not c2:
            rs.text_resume = summary
            rs.save()

    # Marquer la réunion comme "terminer" après sauvegarde (optionnel)
    if reunion.status != "terminer":
        reunion.status = "terminer"
        reunion.save(update_fields=["status"])

    return JsonResponse({"ok": True, "transcription_id": tr.id})


# =============================================================================
# Rapport (DOCX) (utilisateur uniquement)
# =============================================================================
@login_required
def meeting_report_view(request, pk):
    """Aperçu d’un rapport (utilisateur uniquement)."""
    if forbid_admin_on_meetings(request):
        return HttpResponseForbidden("Accès refusé.")

    r = get_object_or_404(Reunion.objects.select_related("utilisateur"), pk=pk)
    if not _can_view_meeting(r, request.user):
        return HttpResponseForbidden("Accès refusé.")

    transcription = getattr(r, "transcription", None)
    resume = getattr(transcription, "resume", None) if transcription else None
    audios = r.audios.all() if hasattr(r, "audios") else []
    total_seconds = sum(getattr(a, "duree", 0) for a in audios)
    total_hours = round(total_seconds / 3600, 2)

    decisions = getattr(r, "decisions", None)  # si vous avez un modèle Decision

    ctx = {
        "reunion": r,
        "participants": r.participants.all(),
        "transcription": transcription,
        "resume": resume,
        "decisions": decisions,
        "audios": audios,
        "total_audio_hours": total_hours,
        "now": timezone.now(),
        "generate_url": reverse("generate_report", args=[r.id]),
    }
    return render(request, "HTML/meeting_report.html", ctx)


@login_required
def generate_report(request, reunion_id: int):
    """
    Génère un DOCX avec infos réunion + résumé + transcription.
    Nécessite: pip install python-docx
    """
    if forbid_admin_on_meetings(request):
        return HttpResponseForbidden("Accès refusé.")

    r = get_object_or_404(Reunion.objects.select_related("utilisateur"), pk=reunion_id)
    if not _can_view_meeting(r, request.user):
        return HttpResponseForbidden("Accès refusé.")

    tr = getattr(r, "transcription", None)
    if not tr or not (tr.text_transcrit or "").strip():
        messages.error(request, "Aucune transcription enregistrée pour cette réunion.")
        return redirect("transcription", r.id)

    rs = getattr(tr, "resume", None)
    if not rs:
        # Crée un résumé minimal si absent (OneToOne avec Rapport)
        rs = Resume.objects.create(
            transcription=tr,
            text_resume="(Résumé non fourni — à compléter)",
            langue=(getattr(tr, "langue", "fr") or "fr"),
        )

    # ------- DOCX -------
    doc = Document()
    doc.add_heading("Rapport de réunion", level=1)

    # Logo optionnel
    logo_path = os.path.join(getattr(settings, "MEDIA_ROOT", ""), "logo.png")
    if os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(1.2))
        except Exception:
            pass

    # Métadonnées
    doc.add_paragraph(f"Titre : {r.titre}")
    doc.add_paragraph(f"Date : {r.date_r.strftime('%d/%m/%Y')} à {r.heure_r.strftime('%H:%M')}")
    doc.add_paragraph(f"Statut : {r.get_status_display()}")

    # Participants
    if r.participants.exists():
        doc.add_heading("Participants", level=2)
        for u in r.participants.all():
            doc.add_paragraph(u.get_full_name() or u.username, style="List Bullet")

    # Résumé
    doc.add_heading("Résumé", level=2)
    doc.add_paragraph(rs.text_resume or "(non disponible)")

    # Décisions (si vous avez de la donnée)
    decisions = getattr(r, "decisions", None)
    if decisions:
        doc.add_heading("Décisions", level=2)
        for d in decisions.all() if hasattr(decisions, "all") else decisions:
            doc.add_paragraph(f"• {getattr(d, 'titre', str(d))}", style="List Bullet")

    # Transcription
    doc.add_heading("Transcription", level=2)
    doc.add_paragraph(tr.text_transcrit)

    # ------- Sauvegarde et attachement -------
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp.flush()
        tmp_path = tmp.name

    rapport, _ = Rapport.objects.get_or_create(resume=rs)
    filename = f"rapport_reunion_{r.id}_{timezone.now().strftime('%Y%m%d_%H%M')}.docx"
    with open(tmp_path, "rb") as f:
        rapport.fichier.save(filename, f, save=True)

    os.remove(tmp_path)

    # Téléchargement direct
    return FileResponse(rapport.fichier.open("rb"), as_attachment=True, filename=filename)
